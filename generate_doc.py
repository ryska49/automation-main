from docx import Document
from docx.shared import Inches
import os

# Paths for your code and output image folders
code_folder = "/home/kartik/Desktop/24mcf1r19_kartik/codes"
output_folder = "/home/kartik/Desktop/24mcf1r19_kartik/outputs"

# Create a new Word Document
doc = Document()
doc.add_heading('Code and Output Document', level=1)

# Iterate through each code file
for filename in sorted(os.listdir(code_folder)):
    if filename.endswith('.cpp'):
        # Add code file name as heading
        doc.add_heading(f'{filename}', level=2)
        
        # Add the code file content
        with open(os.path.join(code_folder, filename), 'r') as file:
            code_content = file.read()
            doc.add_paragraph(code_content)

        # Add corresponding output image
        image_filename = f'OUTPUT_{filename[1]}.png'
        image_path = os.path.join(output_folder, image_filename)
        if os.path.exists(image_path):
            doc.add_heading(f'Output for {filename}', level=3)
            doc.add_picture(image_path, width=Inches(5.0))

# Save the Word Document
doc.save('/home/kartik/Desktop/Code_and_Output_Document.docx')
print("Document created successfully!")
